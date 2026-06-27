import io

from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import mm
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.platypus.flowables import HRFlowable

from services.gemini_service import GeminiService

SYSTEM_INSTRUCTION = """
Você é um redator especialista em currículos otimizados para ATS (Applicant Tracking System).
Gere conteúdo claro, objetivo, com verbos de ação e palavras-chave da vaga.
"""


class ResumeService:
    def __init__(self):
        self.gemini = GeminiService()

    def generate_resume_content(self, candidate_profile: dict, job_requirements: dict) -> dict:
        prompt = f"""
        Gere o conteúdo de um currículo otimizado para ATS em JSON com a estrutura:
        {{
          "dados_pessoais": {{}},
          "resumo_profissional": "",
          "formacao": [],
          "experiencia": [],
          "projetos": [],
          "certificacoes": [],
          "competencias": [],
          "idiomas": []
        }}

        Perfil do candidato: {candidate_profile}
        Requisitos da vaga: {job_requirements}
        """
        return self.gemini.generate_json(prompt, system_instruction=SYSTEM_INSTRUCTION)

    def render_pdf(self, resume_content: dict) -> bytes:
        buf = io.BytesIO()

        styles = getSampleStyleSheet()

        title_style = ParagraphStyle(
            "ResumeTitle",
            parent=styles["Heading1"],
            fontSize=20,
            leading=24,
            spaceAfter=4,
            textColor="#1e3a8a",
            fontName="Helvetica-Bold",
        )
        heading_style = ParagraphStyle(
            "ResumeHeading",
            parent=styles["Heading2"],
            fontSize=13,
            leading=16,
            spaceBefore=14,
            spaceAfter=6,
            textColor="#2563eb",
            fontName="Helvetica-Bold",
        )
        normal_style = ParagraphStyle(
            "ResumeNormal",
            parent=styles["Normal"],
            fontSize=10,
            leading=13,
            spaceAfter=4,
            fontName="Helvetica",
        )
        small_style = ParagraphStyle(
            "ResumeSmall",
            parent=styles["Normal"],
            fontSize=9,
            leading=11,
            textColor="#64748b",
            fontName="Helvetica",
        )

        doc = SimpleDocTemplate(
            buf,
            pagesize=A4,
            leftMargin=20 * mm,
            rightMargin=20 * mm,
            topMargin=20 * mm,
            bottomMargin=20 * mm,
        )

        elements = []

        dp = resume_content.get("dados_pessoais") or {}
        nome = dp.get("nome", "Candidato")
        email = dp.get("email", "")
        telefone = dp.get("telefone", "")
        contato = " | ".join(filter(None, [email, telefone]))
        elements.append(Paragraph(nome, title_style))
        if contato:
            elements.append(Paragraph(contato, small_style))
        elements.append(HRFlowable(width="100%", color="#e2e8f0", thickness=1))

        resumo = resume_content.get("resumo_profissional")
        if resumo:
            elements.append(Paragraph("Resumo Profissional", heading_style))
            elements.append(Paragraph(resumo, normal_style))

        experiencia = resume_content.get("experiencia") or []
        if experiencia:
            elements.append(Paragraph("Experiência Profissional", heading_style))
            for exp in experiencia:
                cargo = exp.get("cargo", "")
                empresa = exp.get("empresa", "")
                descricao = exp.get("descricao", "")
                if cargo and empresa:
                    elements.append(Paragraph(f"<b>{cargo}</b> — {empresa}", normal_style))
                elif cargo:
                    elements.append(Paragraph(f"<b>{cargo}</b>", normal_style))
                if descricao:
                    elements.append(Paragraph(descricao, small_style))

        formacao = resume_content.get("formacao") or []
        if formacao:
            elements.append(Paragraph("Formação", heading_style))
            for f in formacao:
                curso = f.get("curso", "")
                instituicao = f.get("instituicao", "")
                texto = curso
                if instituicao:
                    texto += f" — {instituicao}"
                elements.append(Paragraph(texto, normal_style))

        competencias = resume_content.get("competencias") or []
        if competencias:
            elements.append(Paragraph("Competências", heading_style))
            skills_text = " • ".join(competencias)
            elements.append(Paragraph(skills_text, normal_style))

        idiomas = resume_content.get("idiomas") or []
        if idiomas:
            elements.append(Paragraph("Idiomas", heading_style))
            lang_text = " • ".join(idiomas)
            elements.append(Paragraph(lang_text, normal_style))

        certificacoes = resume_content.get("certificacoes") or []
        if certificacoes:
            elements.append(Paragraph("Certificações", heading_style))
            for cert in certificacoes:
                if isinstance(cert, str):
                    elements.append(Paragraph(cert, normal_style))
                elif isinstance(cert, dict):
                    elements.append(Paragraph(cert.get("nome", str(cert)), normal_style))

        projetos = resume_content.get("projetos") or []
        if projetos:
            elements.append(Paragraph("Projetos", heading_style))
            for proj in projetos:
                if isinstance(proj, str):
                    elements.append(Paragraph(proj, normal_style))
                elif isinstance(proj, dict):
                    title = proj.get("titulo") or proj.get("nome", "")
                    desc = proj.get("descricao", "")
                    if title:
                        elements.append(Paragraph(f"<b>{title}</b>", normal_style))
                    if desc:
                        elements.append(Paragraph(desc, small_style))

        doc.build(elements)
        return buf.getvalue()

    def render_docx(self, resume_content: dict) -> bytes:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Calibri"
        font.size = Pt(10)

        dp = resume_content.get("dados_pessoais") or {}
        nome = dp.get("nome", "Candidato")
        email = dp.get("email", "")
        telefone = dp.get("telefone", "")

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(nome)
        run.bold = True
        run.font.size = Pt(20)
        run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)

        if email or telefone:
            p = doc.add_paragraph()
            p.add_run(" | ".join(filter(None, [email, telefone]))).font.size = Pt(9)

        def add_heading(text):
            h = doc.add_heading(text, level=2)
            for run in h.runs:
                run.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)
                run.font.size = Pt(13)

        resumo = resume_content.get("resumo_profissional")
        if resumo:
            add_heading("Resumo Profissional")
            doc.add_paragraph(resumo)

        experiencia = resume_content.get("experiencia") or []
        if experiencia:
            add_heading("Experiência Profissional")
            for exp in experiencia:
                cargo = exp.get("cargo", "")
                empresa = exp.get("empresa", "")
                descricao = exp.get("descricao", "")
                p = doc.add_paragraph()
                if cargo and empresa:
                    run = p.add_run(f"{cargo} — {empresa}")
                elif cargo:
                    run = p.add_run(cargo)
                run.bold = True
                if descricao:
                    doc.add_paragraph(descricao)

        formacao = resume_content.get("formacao") or []
        if formacao:
            add_heading("Formação")
            for f in formacao:
                curso = f.get("curso", "")
                instituicao = f.get("instituicao", "")
                texto = curso
                if instituicao:
                    texto += f" — {instituicao}"
                doc.add_paragraph(texto)

        competencias = resume_content.get("competencias") or []
        if competencias:
            add_heading("Competências")
            doc.add_paragraph(" • ".join(competencias))

        idiomas = resume_content.get("idiomas") or []
        if idiomas:
            add_heading("Idiomas")
            doc.add_paragraph(" • ".join(idiomas))

        certificacoes = resume_content.get("certificacoes") or []
        if certificacoes:
            add_heading("Certificações")
            for cert in certificacoes:
                if isinstance(cert, str):
                    doc.add_paragraph(cert)
                elif isinstance(cert, dict):
                    doc.add_paragraph(cert.get("nome", str(cert)))

        projetos = resume_content.get("projetos") or []
        if projetos:
            add_heading("Projetos")
            for proj in projetos:
                if isinstance(proj, str):
                    doc.add_paragraph(proj)
                elif isinstance(proj, dict):
                    title = proj.get("titulo") or proj.get("nome", "")
                    desc = proj.get("descricao", "")
                    p = doc.add_paragraph()
                    if title:
                        run = p.add_run(title)
                        run.bold = True
                    if desc:
                        doc.add_paragraph(desc)

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()
